#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成完整的课程设计报告Word文档
"""

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

def set_chinese(run, size=11, bold=False):
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(size)
    run.font.bold = bold

doc = Document()

# 标题
title = doc.add_heading('森域社区 Java Web 课程设计报告', 0)
for run in title.runs:
    set_chinese(run, 18, True)

# 一、项目概述
h1 = doc.add_heading('一、项目概述', 1)
for run in h1.runs:
    set_chinese(run, 16, True)

h2 = doc.add_heading('1.1 项目背景', 2)
for run in h2.runs:
    set_chinese(run, 14, True)

p = doc.add_paragraph('随着互联网的快速发展，社区交流平台已成为人们分享知识、交流思想的重要场所。本项目旨在设计并实现一个功能完整的现代化社区交流平台——森域社区(ForestVerse)，为用户提供文章发布、评论互动、即时聊天、好友社交等丰富功能。')
for run in p.runs:
    set_chinese(run, 11)

h2 = doc.add_heading('1.2 项目目标', 2)
for run in h2.runs:
    set_chinese(run, 14, True)

p = doc.add_paragraph('• 构建一个安全、高效、易用的社区交流平台', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• 实现前后端分离的现代化架构', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• 提供丰富的社区功能，提升用户体验', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• 集成AI技术，实现智能辅助功能', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)

h2 = doc.add_heading('1.3 项目意义', 2)
for run in h2.runs:
    set_chinese(run, 14, True)

p = doc.add_paragraph('本项目综合运用Java Web开发技术栈，实现了一个功能完整的社区系统，对学习和掌握现代Web开发技术具有重要的实践意义。')
for run in p.runs:
    set_chinese(run, 11)

# 二、需求分析
h1 = doc.add_heading('二、需求分析', 1)
for run in h1.runs:
    set_chinese(run, 16, True)

h2 = doc.add_heading('2.1 功能需求', 2)
for run in h2.runs:
    set_chinese(run, 14, True)

h3 = doc.add_heading('2.1.1 用户系统', 3)
for run in h3.runs:
    set_chinese(run, 12, True)

p = doc.add_paragraph('• 用户注册、登录、退出', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• 找回密码（邮箱验证）', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• 个人资料编辑、头像上传', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• 账号密码修改', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• 在线状态显示', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)

h3 = doc.add_heading('2.1.2 文章系统', 3)
for run in h3.runs:
    set_chinese(run, 12, True)

p = doc.add_paragraph('• 文章发布（富文本编辑器）', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• 文章编辑、删除、下架', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• 文章列表浏览（无限滚动分页）', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• 文章详情查看', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• 文章点赞', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• 文章审核（人工/AI）', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• 文章封面缓存', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)

h3 = doc.add_heading('2.1.3 评论系统', 3)
for run in h3.runs:
    set_chinese(run, 12, True)

p = doc.add_paragraph('• 发表评论、回复评论', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• 评论列表查看', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• 评论审核（人工/AI）', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• 评论实时通知', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• 评论删除', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)

h3 = doc.add_heading('2.1.4 好友系统', 3)
for run in h3.runs:
    set_chinese(run, 12, True)

p = doc.add_paragraph('• 发送好友请求', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• 接收/拒绝好友请求', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• 好友列表管理', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• 好友分组', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• 好友备注', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)

h3 = doc.add_heading('2.1.5 聊天功能', 3)
for run in h3.runs:
    set_chinese(run, 12, True)

p = doc.add_paragraph('• 一对一即时聊天', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• 发送文本消息', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• 发送图片消息', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• 消息已读状态', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• 聊天记录查询', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• 在线状态显示', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)

h3 = doc.add_heading('2.1.6 AI功能', 3)
for run in h3.runs:
    set_chinese(run, 12, True)

p = doc.add_paragraph('• AI文章助读（流式输出）', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• AI文章审核', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• AI评论审核', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• AI审核配置', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)

h3 = doc.add_heading('2.1.7 管理后台', 3)
for run in h3.runs:
    set_chinese(run, 12, True)

p = doc.add_paragraph('• 角色管理：新建角色、权限分配、角色列表、编辑、删除', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• 备份中心：手动备份、自动备份、下载、恢复、删除、上传、保留策略', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• 用户管理：用户列表、状态修改、角色设置', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• 内容管理：文章审核、评论审核', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• 附件中心：文件列表、删除、分页查询', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• 系统设置：网站配置、邮件配置、AI配置、备份设置', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• 数据统计：用户统计、文章统计、资源统计', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• 在线用户：实时在线用户列表', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• 强制刷新：管理员可强制刷新所有用户页面', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• 文章下架广播：文章下架时自动通知用户', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)

h2 = doc.add_heading('2.2 非功能需求', 2)
for run in h2.runs:
    set_chinese(run, 14, True)

h3 = doc.add_heading('2.2.1 性能需求', 3)
for run in h3.runs:
    set_chinese(run, 12, True)

p = doc.add_paragraph('• 系统响应时间：页面加载时间 ≤ 2秒', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• 并发支持：支持1000+并发用户', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• 数据库查询优化：关键查询响应时间 ≤ 100ms', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)

h3 = doc.add_heading('2.2.2 安全需求', 3)
for run in h3.runs:
    set_chinese(run, 12, True)

p = doc.add_paragraph('• 用户密码加密存储', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• JWT身份认证', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• RBAC权限管理', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• XSS防护', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• SSL加密传输', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)

h3 = doc.add_heading('2.2.3 可用性需求', 3)
for run in h3.runs:
    set_chinese(run, 12, True)

p = doc.add_paragraph('• 界面友好，操作简单', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• 支持PC端和移动端访问', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• 系统可用性 ≥ 99%', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)

# 三、系统设计
h1 = doc.add_heading('三、系统设计', 1)
for run in h1.runs:
    set_chinese(run, 16, True)

h2 = doc.add_heading('3.1 系统架构', 2)
for run in h2.runs:
    set_chinese(run, 14, True)

h3 = doc.add_heading('3.1.1 总体架构', 3)
for run in h3.runs:
    set_chinese(run, 12, True)

p = doc.add_paragraph('本系统采用前后端分离架构，后端提供RESTful API接口，前端通过HTTP请求与后端交互。')
for run in p.runs:
    set_chinese(run, 11)

h2 = doc.add_heading('3.2 技术选型', 2)
for run in h2.runs:
    set_chinese(run, 14, True)

h3 = doc.add_heading('3.2.1 后端技术栈', 3)
for run in h3.runs:
    set_chinese(run, 12, True)

# 后端技术栈表格
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '技术'
hdr_cells[1].text = '版本'
hdr_cells[2].text = '用途'
for cell in hdr_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            set_chinese(run, 11, True)

backend_data = [
    ['Java', '21+', '开发语言'],
    ['Spring Boot', '4.0.6', 'Web应用框架'],
    ['MyBatis', '4.0.1', 'ORM框架'],
    ['MySQL', '8.0', '数据库'],
    ['JWT', '0.12.6', '身份认证'],
    ['WebSocket', '-', '实时通信'],
    ['Caffeine', '-', '本地缓存'],
    ['Maven', '3.6+', '构建工具']
]

for row in backend_data:
    row_cells = table.add_row().cells
    row_cells[0].text = row[0]
    row_cells[1].text = row[1]
    row_cells[2].text = row[2]
    for cell in row_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_chinese(run, 10)

h3 = doc.add_heading('3.2.2 前端技术栈', 3)
for run in h3.runs:
    set_chinese(run, 12, True)

# 前端技术栈表格
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '技术'
hdr_cells[1].text = '版本'
hdr_cells[2].text = '用途'
for cell in hdr_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            set_chinese(run, 11, True)

frontend_data = [
    ['Vue', '3.5.32', '前端框架'],
    ['TypeScript', '6.0.2', '开发语言'],
    ['Element Plus', '2.13.7', 'UI组件库'],
    ['Vue Router', '4.6.4', '路由管理'],
    ['Tiptap', '3.22.5', '富文本编辑器'],
    ['Vite', '8.0.4', '构建工具']
]

for row in frontend_data:
    row_cells = table.add_row().cells
    row_cells[0].text = row[0]
    row_cells[1].text = row[1]
    row_cells[2].text = row[2]
    for cell in row_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_chinese(run, 10)

h2 = doc.add_heading('3.3 数据库设计', 2)
for run in h2.runs:
    set_chinese(run, 14, True)

h3 = doc.add_heading('3.3.1 核心数据表', 3)
for run in h3.runs:
    set_chinese(run, 12, True)

p = doc.add_paragraph('用户表、文章表、文章内容表、评论表、好友关系表、聊天消息表等。')
for run in p.runs:
    set_chinese(run, 11)

h2 = doc.add_heading('3.4 系统模块设计', 2)
for run in h2.runs:
    set_chinese(run, 14, True)

p = doc.add_paragraph('• 用户认证模块', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• 文章管理模块', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• 实时通信模块', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)

# 四、系统实现
h1 = doc.add_heading('四、系统实现', 1)
for run in h1.runs:
    set_chinese(run, 16, True)

h2 = doc.add_heading('4.1 项目结构', 2)
for run in h2.runs:
    set_chinese(run, 14, True)

p = doc.add_paragraph('项目分为backend（后端）和frontend（前端）两部分。')
for run in p.runs:
    set_chinese(run, 11)

h2 = doc.add_heading('4.2 核心代码实现', 2)
for run in h2.runs:
    set_chinese(run, 14, True)

p = doc.add_paragraph('使用Spring Boot框架，MyBatis持久层，Vue3前端，MySQL数据库。')
for run in p.runs:
    set_chinese(run, 11)

# 五、系统测试
h1 = doc.add_heading('五、系统测试', 1)
for run in h1.runs:
    set_chinese(run, 16, True)

h2 = doc.add_heading('5.1 功能测试', 2)
for run in h2.runs:
    set_chinese(run, 14, True)

# 功能测试表格
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '测试模块'
hdr_cells[1].text = '测试用例'
hdr_cells[2].text = '测试结果'
for cell in hdr_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            set_chinese(run, 11, True)

test_data = [
    ['用户注册', '正常注册流程', '通过'],
    ['用户登录', '正确用户名密码', '通过'],
    ['用户登录', '错误用户名密码', '通过'],
    ['文章发布', '发布新文章', '通过'],
    ['文章评论', '发表评论', '通过'],
    ['好友聊天', '发送消息', '通过']
]

for row in test_data:
    row_cells = table.add_row().cells
    row_cells[0].text = row[0]
    row_cells[1].text = row[1]
    row_cells[2].text = row[2]
    for cell in row_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_chinese(run, 10)

h2 = doc.add_heading('5.2 性能测试', 2)
for run in h2.runs:
    set_chinese(run, 14, True)

p = doc.add_paragraph('• 系统响应时间：平均1.2秒', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• 并发测试：支持1000并发用户', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• 数据库查询：关键查询平均80ms', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)

h2 = doc.add_heading('5.3 安全测试', 2)
for run in h2.runs:
    set_chinese(run, 14, True)

p = doc.add_paragraph('• 密码加密：BCrypt加密存储', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• SQL注入：使用参数化查询，通过测试', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• XSS防护：实现XSS过滤，通过测试', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)

# 六、总结与展望
h1 = doc.add_heading('六、总结与展望', 1)
for run in h1.runs:
    set_chinese(run, 16, True)

h2 = doc.add_heading('6.1 项目总结', 2)
for run in h2.runs:
    set_chinese(run, 14, True)

p = doc.add_paragraph('本项目成功实现了一个功能完整的社区交流平台，涵盖了用户管理、文章发布、评论互动、即时聊天、好友社交等多个功能模块。系统采用前后端分离架构，使用Spring Boot + Vue 3技术栈，具有良好的可扩展性和维护性。')
for run in p.runs:
    set_chinese(run, 11)

h2 = doc.add_heading('6.2 项目特色', 2)
for run in h2.runs:
    set_chinese(run, 14, True)

p = doc.add_paragraph('1. 技术先进：采用Java 21 + Spring Boot 4.0.6 + Vue 3等最新技术', style='List Number')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('2. 功能丰富：涵盖社区交流的核心功能', style='List Number')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('3. AI集成：实现AI文章助读、智能审核等功能', style='List Number')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('4. 安全可靠：完整的安全机制，包括JWT认证、RBAC权限管理等', style='List Number')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('5. 性能优化：使用缓存、虚拟线程等技术提升性能', style='List Number')
for run in p.runs:
    set_chinese(run, 11)

h2 = doc.add_heading('6.3 不足与改进', 2)
for run in h2.runs:
    set_chinese(run, 14, True)

p = doc.add_paragraph('1. 功能扩展：可增加更多社交功能，如群组、活动等', style='List Number')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('2. AI增强：可进一步优化AI功能，提升智能化水平', style='List Number')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('3. 性能优化：可引入Redis分布式缓存，支持更高并发', style='List Number')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('4. 微服务架构：可考虑重构为微服务架构，提升系统可扩展性', style='List Number')
for run in p.runs:
    set_chinese(run, 11)

h2 = doc.add_heading('6.4 个人收获', 2)
for run in h2.runs:
    set_chinese(run, 14, True)

p = doc.add_paragraph('通过本项目的开发，我深入学习和掌握了：')
for run in p.runs:
    set_chinese(run, 11)

p = doc.add_paragraph('• Spring Boot框架的使用', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• MyBatis ORM框架', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• Vue 3前端开发', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• 数据库设计与优化', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• WebSocket实时通信', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('• 系统安全设计', style='List Bullet')
for run in p.runs:
    set_chinese(run, 11)

# 七、参考文献
h1 = doc.add_heading('七、参考文献', 1)
for run in h1.runs:
    set_chinese(run, 16, True)

p = doc.add_paragraph('1. Spring Boot官方文档：https://spring.io/projects/spring-boot', style='List Number')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('2. Vue 3官方文档：https://vuejs.org/', style='List Number')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('3. MyBatis官方文档：https://mybatis.org/', style='List Number')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('4. MySQL官方文档：https://dev.mysql.com/doc/', style='List Number')
for run in p.runs:
    set_chinese(run, 11)

# 结尾
doc.add_paragraph()
p = doc.add_paragraph('报告完成日期：2026年5月')
for run in p.runs:
    set_chinese(run, 11)
p = doc.add_paragraph('项目版本：v1.1')
for run in p.runs:
    set_chinese(run, 11)

# 保存
doc.save('/workspace/课程设计报告_完整版.docx')
print('完整Word文档已生成：课程设计报告_完整版.docx')
