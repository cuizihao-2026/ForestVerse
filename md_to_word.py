#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将课程设计报告Markdown转换为Word文档
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

def set_chinese_font(run):
    """设置中文字体"""
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_heading(doc, text, level):
    """添加标题"""
    heading = doc.add_heading(text, level)
    for run in heading.runs:
        set_chinese_font(run)
        if level == 0:
            run.font.size = Pt(18)
            run.font.bold = True
        elif level == 1:
            run.font.size = Pt(16)
            run.font.bold = True
        elif level == 2:
            run.font.size = Pt(14)
            run.font.bold = True
        elif level == 3:
            run.font.size = Pt(12)
            run.font.bold = True
    return heading

def add_paragraph(doc, text, bold=False, italic=False):
    """添加段落"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_chinese_font(run)
    run.font.size = Pt(11)
    run.font.bold = bold
    run.font.italic = italic
    return p

def add_list_item(doc, text, level=0):
    """添加列表项"""
    p = doc.add_paragraph(text, style='List Bullet' if level == 0 else 'List Bullet 2')
    p.paragraph_format.left_indent = Inches(0.25 * level)
    for run in p.runs:
        set_chinese_font(run)
        run.font.size = Pt(11)
    return p

def add_table(doc, headers, rows):
    """添加表格"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # 添加表头
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                set_chinese_font(run)
                run.font.size = Pt(11)
                run.font.bold = True
    
    # 添加数据行
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    set_chinese_font(run)
                    run.font.size = Pt(10)
    
    return table

def add_code_block(doc, code):
    """添加代码块"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    return p

def parse_markdown_to_word(md_file, docx_file):
    """解析Markdown并生成Word文档"""
    doc = Document()
    
    # 设置默认样式
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.styles['Normal'].font.size = Pt(11)
    
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    in_code_block = False
    code_content = []
    in_table = False
    table_headers = []
    table_rows = []
    
    while i < len(lines):
        line = lines[i].rstrip('\n')
        
        # 代码块处理
        if line.startswith('```'):
            if in_code_block:
                # 结束代码块
                add_code_block(doc, '\n'.join(code_content))
                code_content = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_content.append(line)
            i += 1
            continue
        
        # 空行
        if not line.strip():
            i += 1
            continue
        
        # 标题处理
        if line.startswith('#'):
            level = len(line.split()[0])
            text = line.lstrip('#').strip()
            add_heading(doc, text, level - 1)
            i += 1
            continue
        
        # 分隔线
        if line.startswith('---'):
            i += 1
            continue
        
        # 列表项
        if line.startswith('- '):
            text = line[2:].strip()
            add_list_item(doc, text)
            i += 1
            continue
        
        # 表格处理
        if line.startswith('|'):
            if not in_table:
                in_table = True
                table_headers = [cell.strip() for cell in line.split('|')[1:-1]]
                i += 1
                # 跳过表格分隔线
                if i < len(lines) and lines[i].startswith('|'):
                    i += 1
            else:
                # 检查是否还是表格行
                if line.startswith('|'):
                    row = [cell.strip() for cell in line.split('|')[1:-1]]
                    table_rows.append(row)
                    i += 1
                else:
                    # 表格结束
                    if table_headers and table_rows:
                        add_table(doc, table_headers, table_rows)
                    in_table = False
                    table_headers = []
                    table_rows = []
            continue
        
        # 如果之前在表格中，现在结束了
        if in_table and table_headers and table_rows:
            add_table(doc, table_headers, table_rows)
            in_table = False
            table_headers = []
            table_rows = []
        
        # 普通段落
        add_paragraph(doc, line)
        i += 1
    
    # 如果最后还有未处理的表格
    if in_table and table_headers and table_rows:
        add_table(doc, table_headers, table_rows)
    
    # 保存文档
    doc.save(docx_file)
    print(f"Word文档已生成: {docx_file}")

if __name__ == "__main__":
    parse_markdown_to_word('/workspace/课程设计报告.md', '/workspace/课程设计报告.docx')
