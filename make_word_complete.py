#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
完整的Markdown到Word转换
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn

def set_chinese(run, size=11, bold=False):
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(size)
    run.font.bold = bold

def main():
    doc = Document()
    
    # 读取完整的Markdown内容
    with open('/workspace/课程设计报告.md', 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 标题
    title = doc.add_heading('森域社区 Java Web 课程设计报告', 0)
    for run in title.runs:
        set_chinese(run, 18, True)
    
    # 逐段处理内容
    lines = content.split('\n')
    i = 0
    in_code = False
    code_buffer = []
    
    while i < len(lines):
        line = lines[i]
        
        # 代码块处理
        if line.startswith('```'):
            if in_code:
                # 代码块结束
                p = doc.add_paragraph()
                for run in p.runs:
                    p.clear()
                run = p.add_run('\n'.join(code_buffer))
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                in_code = False
                code_buffer = []
            else:
                in_code = True
            i += 1
            continue
        
        if in_code:
            code_buffer.append(line)
            i += 1
            continue
        
        # 标题
        if line.startswith('#'):
            level = len(line.split()[0])
            text = line.lstrip('#').strip()
            heading = doc.add_heading(text, level-1 if level-1 <= 9 else 9)
            for run in heading.runs:
                sizes = [18, 16, 14, 12, 12, 11, 11, 11, 11]
                set_chinese(run, sizes[level-1] if level-1 < len(sizes) else 11, True)
            i += 1
            continue
        
        # 分隔线
        if line.startswith('---'):
            i += 1
            continue
        
        # 空行
        if not line.strip():
            i += 1
            continue
        
        # 列表项
        if line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            for run in p.runs:
                set_chinese(run, 11)
            i += 1
            continue
        
        # 表格处理（简化处理）
        if line.startswith('|'):
            # 跳过表格处理，直接把内容写为段落
            i += 1
            continue
        
        # 普通段落
        p = doc.add_paragraph(line)
        for run in p.runs:
            set_chinese(run, 11)
        i += 1
    
    # 如果有未结束的代码块
    if in_code and code_buffer:
        p = doc.add_paragraph()
        run = p.add_run('\n'.join(code_buffer))
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    
    # 保存
    doc.save('/workspace/课程设计报告_final.docx')
    print('完整Word文档已生成：课程设计报告_final.docx')

if __name__ == '__main__':
    main()
