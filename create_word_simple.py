#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
简单直接的Word文档生成
"""

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

def set_font(run):
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(11)

doc = Document()

# 标题
title = doc.add_heading('森域社区 Java Web 课程设计报告', 0)
for run in title.runs:
    set_font(run)
    run.font.size = Pt(18)
    run.font.bold = True

# 一、项目概述
h1 = doc.add_heading('一、项目概述', 1)
for run in h1.runs:
    set_font(run)
    run.font.size = Pt(16)

h2 = doc.add_heading('1.1 项目背景', 2)
for run in h2.runs:
    set_font(run)
    run.font.size = Pt(14)

p = doc.add_paragraph('随着互联网的快速发展，社区交流平台已成为人们分享知识、交流思想的重要场所。本项目旨在设计并实现一个功能完整的现代化社区交流平台——森域社区(ForestVerse)，为用户提供文章发布、评论互动、即时聊天、好友社交等丰富功能。')
for run in p.runs:
    set_font(run)

h2 = doc.add_heading('1.2 项目目标', 2)
for run in h2.runs:
    set_font(run)
    run.font.size = Pt(14)

p = doc.add_paragraph('• 构建一个安全、高效、易用的社区交流平台', style='List Bullet')
for run in p.runs:
    set_font(run)
p = doc.add_paragraph('• 实现前后端分离的现代化架构', style='List Bullet')
for run in p.runs:
    set_font(run)
p = doc.add_paragraph('• 提供丰富的社区功能，提升用户体验', style='List Bullet')
for run in p.runs:
    set_font(run)
p = doc.add_paragraph('• 集成AI技术，实现智能辅助功能', style='List Bullet')
for run in p.runs:
    set_font(run)

h2 = doc.add_heading('1.3 项目意义', 2)
for run in h2.runs:
    set_font(run)
    run.font.size = Pt(14)

p = doc.add_paragraph('本项目综合运用Java Web开发技术栈，实现了一个功能完整的社区系统，对学习和掌握现代Web开发技术具有重要的实践意义。')
for run in p.runs:
    set_font(run)

# 二、需求分析
h1 = doc.add_heading('二、需求分析', 1)
for run in h1.runs:
    set_font(run)
    run.font.size = Pt(16)

h2 = doc.add_heading('2.1 功能需求', 2)
for run in h2.runs:
    set_font(run)
    run.font.size = Pt(14)

h3 = doc.add_heading('2.1.1 用户系统', 3)
for run in h3.runs:
    set_font(run)
    run.font.size = Pt(12)

p = doc.add_paragraph('• 用户注册、登录、退出', style='List Bullet')
for run in p.runs:
    set_font(run)
p = doc.add_paragraph('• 找回密码（邮箱验证）', style='List Bullet')
for run in p.runs:
    set_font(run)
p = doc.add_paragraph('• 个人资料编辑、头像上传', style='List Bullet')
for run in p.runs:
    set_font(run)
p = doc.add_paragraph('• 账号密码修改', style='List Bullet')
for run in p.runs:
    set_font(run)
p = doc.add_paragraph('• 在线状态显示', style='List Bullet')
for run in p.runs:
    set_font(run)

h3 = doc.add_heading('2.1.2 文章系统', 3)
for run in h3.runs:
    set_font(run)
    run.font.size = Pt(12)

p = doc.add_paragraph('• 文章发布（富文本编辑器）', style='List Bullet')
for run in p.runs:
    set_font(run)
p = doc.add_paragraph('• 文章编辑、删除、下架', style='List Bullet')
for run in p.runs:
    set_font(run)
p = doc.add_paragraph('• 文章列表浏览（无限滚动分页）', style='List Bullet')
for run in p.runs:
    set_font(run)
p = doc.add_paragraph('• 文章详情查看', style='List Bullet')
for run in p.runs:
    set_font(run)
p = doc.add_paragraph('• 文章点赞', style='List Bullet')
for run in p.runs:
    set_font(run)
p = doc.add_paragraph('• 文章审核（人工/AI）', style='List Bullet')
for run in p.runs:
    set_font(run)
p = doc.add_paragraph('• 文章封面缓存', style='List Bullet')
for run in p.runs:
    set_font(run)

# 添加更多章节的简单版本
h1 = doc.add_heading('三、系统设计', 1)
for run in h1.runs:
    set_font(run)
    run.font.size = Pt(16)

h1 = doc.add_heading('四、系统实现', 1)
for run in h1.runs:
    set_font(run)
    run.font.size = Pt(16)

h1 = doc.add_heading('五、系统测试', 1)
for run in h1.runs:
    set_font(run)
    run.font.size = Pt(16)

h1 = doc.add_heading('六、总结与展望', 1)
for run in h1.runs:
    set_font(run)
    run.font.size = Pt(16)

h1 = doc.add_heading('七、参考文献', 1)
for run in h1.runs:
    set_font(run)
    run.font.size = Pt(16)

# 添加表格示例
doc.add_paragraph('\n')
h2 = doc.add_heading('技术栈示例', 2)
for run in h2.runs:
    set_font(run)

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '技术'
hdr_cells[1].text = '版本'
hdr_cells[2].text = '用途'
for cell in hdr_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            set_font(run)
            run.font.bold = True

data = [
    ['Java', '21+', '开发语言'],
    ['Spring Boot', '4.0.6', 'Web应用框架'],
    ['Vue', '3.5.32', '前端框架'],
    ['MySQL', '8.0', '数据库']
]

for row in data:
    row_cells = table.add_row().cells
    row_cells[0].text = row[0]
    row_cells[1].text = row[1]
    row_cells[2].text = row[2]
    for cell in row_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_font(run)

doc.add_paragraph('\n')
p = doc.add_paragraph('报告完成日期：2026年5月')
for run in p.runs:
    set_font(run)
p = doc.add_paragraph('项目版本：v1.1')
for run in p.runs:
    set_font(run)

# 保存
doc.save('/workspace/课程设计报告_simple.docx')
print('Word文档已生成：课程设计报告_simple.docx')
